# app.py - Comprehensive Word Formatting Checker with Progress Tracking
from flask import Flask, request, jsonify, render_template_string
from flask_cors import CORS
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_UNDERLINE
import io
import os
import re
from datetime import datetime, timedelta

app = Flask(__name__)
CORS(app)

# Azure Storage Configuration - Using Environment Variables (Secure)
AZURE_STORAGE_CONNECTION_STRING = os.getenv('AZURE_CONNECTION_STRING')
AZURE_ACCOUNT_NAME = os.getenv('AZURE_ACCOUNT_NAME') 
AZURE_ACCOUNT_KEY = os.getenv('AZURE_ACCOUNT_KEY')
CONTAINER_NAME = "practice-files"

# Check if credentials are loaded
if not all([AZURE_STORAGE_CONNECTION_STRING, AZURE_ACCOUNT_NAME, AZURE_ACCOUNT_KEY]):
    print("❌ Error: Azure credentials not found in environment variables")
    print("Run the setup commands in terminal first!")

# Initialize Azure Blob client
blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)

# Define all formatting terms to check
FORMATTING_TERMS = [
    {"searchedWord": "Bold1", "formatCheck": "bold"},
    {"searchedWord": "Italic1", "formatCheck": "italic"},
    {"searchedWord": "Underline1", "formatCheck": "underline"},
    {"searchedWord": "Underline2 Double", "formatCheck": "underlineDouble"},
    {"searchedWord": "Underline3 Dotted", "formatCheck": "underlineDotted"},
    {"searchedWord": "Underline4 Red", "formatCheck": "underlineRed"},
    {"searchedWord": "Superscript", "formatCheck": "superscript"},
    {"searchedWord": "Subscript", "formatCheck": "subscript"},
    {"searchedWord": "Strikethrough1", "formatCheck": "strikethrough"},
    {"searchedWord": "Strikethrough2 Double", "formatCheck": "strikethroughDouble"},
    {"searchedWord": "SMALL CAPS", "formatCheck": "smallCaps"},
    {"searchedWord": "10 point", "formatCheck": "fontSize10"},
    {"searchedWord": "12 point", "formatCheck": "fontSize12"},
    {"searchedWord": "15 point", "formatCheck": "fontSize15"},
    {"searchedWord": "26 point", "formatCheck": "fontSize26"},
    {"searchedWord": "Arial", "formatCheck": "fontArial"},
    {"searchedWord": "Bookman Old Style", "formatCheck": "fontBookman"},
    {"searchedWord": "Comic Sans Ms", "formatCheck": "fontComicSans"},
    {"searchedWord": "Impact", "formatCheck": "fontImpact"},
    {"searchedWord": "Tahoma", "formatCheck": "fontTahoma"},
    {"searchedWord": "Verdana", "formatCheck": "fontVerdana"},
    {"searchedWord": "Font Color Green", "formatCheck": "colorGreen"},
    {"searchedWord": "Font Color Dark Blue", "formatCheck": "colorDarkBlue"},
    {"searchedWord": "Highlight Turquoise", "formatCheck": "highlightTurquoise"},
    {"searchedWord": "Character Spacing Expanded at 5", "formatCheck": "spacingExpanded5"},
    {"searchedWord": "Character Spacing Expanded at 10", "formatCheck": "spacingExpanded10"},
    {"searchedWord": "Character Spacing Condensed at 2", "formatCheck": "spacingCondensed2"},
    {"searchedWord": "Align Right", "formatCheck": "alignRight"},
    {"searchedWord": "Align Center", "formatCheck": "alignCenter"},
    {"searchedWord": "Align Left", "formatCheck": "alignLeft"}
]

def rgb_to_hex(rgb_color):
    """Convert RGBColor to hex string"""
    if rgb_color is None:
        return None
    try:
        return f"#{rgb_color.r:02x}{rgb_color.g:02x}{rgb_color.b:02x}"
    except:
        return None

def is_color_in_range(rgb_color, target_color_name):
    """Check if RGB color matches target color range"""
    if rgb_color is None:
        return False
    
    try:
        r, g, b = rgb_color.r, rgb_color.g, rgb_color.b
    except:
        return False
    
    target_color_name = target_color_name.lower()
    
    if target_color_name == 'green':
        return g > 100 and r < 150 and b < 150
    elif target_color_name == 'darkblue':
        return b > 50 and r < 80 and g < 80 and b > r and b > g
    elif target_color_name == 'red':
        return r > 150 and g < 100 and b < 100
    elif target_color_name == 'blue':
        return b > 100 and r < 150 and g < 150
    elif target_color_name == 'turquoise':
        return g > 150 and b > 150 and r < 120
    
    return False

def check_specific_formatting(term, run, paragraph):
    """Check specific formatting for a term"""
    format_check = term["formatCheck"]
    
    try:
        if format_check == "bold":
            is_correct = bool(run.bold) if run.bold is not None else False
            return {
                "correct": is_correct,
                "message": "Bold formatting correct" if is_correct else "Bold formatting missing",
                "debug": f"run.bold = {run.bold}"
            }
        
        elif format_check == "italic":
            is_correct = bool(run.italic) if run.italic is not None else False
            return {
                "correct": is_correct,
                "message": "Italic formatting correct" if is_correct else "Italic formatting missing",
                "debug": f"run.italic = {run.italic}"
            }
        
        elif format_check == "underline":
            is_correct = run.underline is not None and run.underline != WD_UNDERLINE.NONE
            return {
                "correct": is_correct,
                "message": "Underline formatting correct" if is_correct else "Underline formatting missing",
                "debug": f"run.underline = {run.underline}"
            }
        
        elif format_check == "underlineDouble":
            is_correct = run.underline == WD_UNDERLINE.DOUBLE
            return {
                "correct": is_correct,
                "message": "Double underline correct" if is_correct else f"Underline is {run.underline}, should be double",
                "debug": f"run.underline = {run.underline}"
            }
        
        elif format_check == "superscript":
            is_correct = bool(run.font.superscript) if run.font.superscript is not None else False
            return {
                "correct": is_correct,
                "message": "Superscript formatting correct" if is_correct else "Superscript formatting missing",
                "debug": f"run.font.superscript = {run.font.superscript}"
            }
        
        elif format_check == "subscript":
            is_correct = bool(run.font.subscript) if run.font.subscript is not None else False
            return {
                "correct": is_correct,
                "message": "Subscript formatting correct" if is_correct else "Subscript formatting missing",
                "debug": f"run.font.subscript = {run.font.subscript}"
            }
        
        elif format_check == "strikethrough":
            is_correct = bool(run.font.strike) if run.font.strike is not None else False
            return {
                "correct": is_correct,
                "message": "Strikethrough formatting correct" if is_correct else "Strikethrough formatting missing",
                "debug": f"run.font.strike = {run.font.strike}"
            }
        
        elif format_check == "smallCaps":
            is_correct = bool(run.font.small_caps) if run.font.small_caps is not None else False
            return {
                "correct": is_correct,
                "message": "Small caps formatting correct" if is_correct else "Small caps formatting missing",
                "debug": f"run.font.small_caps = {run.font.small_caps}"
            }
        
        elif format_check.startswith("fontSize"):
            target_size = int(format_check.replace("fontSize", ""))
            actual_size = run.font.size.pt if run.font.size else None
            is_correct = actual_size == target_size
            return {
                "correct": is_correct,
                "message": f"{target_size}pt font size correct" if is_correct else f"Font size is {actual_size}pt, should be {target_size}pt",
                "debug": f"run.font.size = {actual_size}pt"
            }
        
        elif format_check.startswith("font"):
            font_name_map = {
                "fontArial": "arial",
                "fontBookman": "bookman",
                "fontComicSans": "comic",
                "fontImpact": "impact",
                "fontTahoma": "tahoma",
                "fontVerdana": "verdana"
            }
            target_font = font_name_map.get(format_check, "")
            actual_font = run.font.name.lower() if run.font.name else ""
            is_correct = target_font in actual_font
            return {
                "correct": is_correct,
                "message": f"{target_font.title()} font correct" if is_correct else f"Font is {run.font.name}, should contain {target_font}",
                "debug": f"run.font.name = {run.font.name}"
            }
        
        elif format_check == "colorGreen":
            is_correct = is_color_in_range(run.font.color.rgb, 'green')
            return {
                "correct": is_correct,
                "message": "Green font color correct" if is_correct else "Font color doesn't match green range",
                "debug": f"font.color.rgb = {rgb_to_hex(run.font.color.rgb) if run.font.color.rgb else 'None'}"
            }
        
        elif format_check == "colorDarkBlue":
            is_correct = is_color_in_range(run.font.color.rgb, 'darkblue')
            return {
                "correct": is_correct,
                "message": "Dark blue font color correct" if is_correct else "Font color doesn't match dark blue range",
                "debug": f"font.color.rgb = {rgb_to_hex(run.font.color.rgb) if run.font.color.rgb else 'None'}"
            }
        
        elif format_check == "highlightTurquoise":
            # Note: python-docx has limited highlight color support
            is_correct = run.font.highlight_color == WD_COLOR_INDEX.TURQUOISE
            return {
                "correct": is_correct,
                "message": "Turquoise highlight correct" if is_correct else "Highlight color doesn't match turquoise",
                "debug": f"font.highlight_color = {run.font.highlight_color}"
            }
        
        elif format_check.startswith("spacing"):
            # Character spacing is not directly supported in python-docx
            return {
                "correct": False,
                "message": "Character spacing check not available in python-docx",
                "debug": "Character spacing requires more advanced docx parsing"
            }
        
        elif format_check.startswith("align"):
            alignment_map = {
                "alignLeft": WD_ALIGN_PARAGRAPH.LEFT,
                "alignCenter": WD_ALIGN_PARAGRAPH.CENTER,
                "alignRight": WD_ALIGN_PARAGRAPH.RIGHT
            }
            target_alignment = alignment_map.get(format_check)
            actual_alignment = paragraph.alignment
            is_correct = actual_alignment == target_alignment
            return {
                "correct": is_correct,
                "message": f"{format_check.replace('align', '')} alignment correct" if is_correct else f"Alignment is {actual_alignment}, should be {target_alignment}",
                "debug": f"paragraph.alignment = {actual_alignment}"
            }
        
        else:
            return {
                "correct": False,
                "message": f"Format check for {format_check} not implemented",
                "debug": f"formatCheck = {format_check}"
            }
    
    except Exception as e:
        return {
            "correct": False,
            "message": f"Error checking {format_check}: {str(e)}",
            "debug": f"Exception: {str(e)}"
        }

def find_text_in_document(doc, search_text):
    """Find all instances of text in document"""
    found_instances = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        paragraph_text = paragraph.text.lower()
        search_lower = search_text.lower()
        
        # Search for exact match
        if search_lower in paragraph_text:
            # Find the specific run containing the text
            for run_idx, run in enumerate(paragraph.runs):
                if search_lower in run.text.lower():
                    found_instances.append({
                        'paragraph': paragraph,
                        'paragraph_idx': para_idx,
                        'run': run,
                        'run_idx': run_idx,
                        'text': run.text
                    })
        
        # Also search for no-space version if original has spaces
        if " " in search_text:
            no_space_search = search_text.replace(" ", "").lower()
            if no_space_search in paragraph_text.replace(" ", ""):
                for run_idx, run in enumerate(paragraph.runs):
                    if no_space_search in run.text.replace(" ", "").lower():
                        found_instances.append({
                            'paragraph': paragraph,
                            'paragraph_idx': para_idx,
                            'run': run,
                            'run_idx': run_idx,
                            'text': run.text
                        })
    
    return found_instances

def check_comprehensive_formatting(doc):
    """Check all formatting terms in document"""
    print("🔍 Starting comprehensive formatting check...")
    
    results = {
        "total_terms": len(FORMATTING_TERMS),
        "found_count": 0,
        "format_correct_count": 0,
        "part1_results": [],
        "part2_results": [],
        "overall_score": 0,
        "debug_info": []
    }
    
    print(f"📊 Total terms to check: {results['total_terms']}")
    print(f"📄 Document has {len(doc.paragraphs)} paragraphs")
    
    found_terms = []
    
    # PART 1: Find all words
    print("🔍 PART 1: Finding words...")
    for i, term in enumerate(FORMATTING_TERMS):
        search_word = term["searchedWord"]
        print(f"  📝 Searching for: '{search_word}' ({i+1}/{len(FORMATTING_TERMS)})")
        
        found_instances = find_text_in_document(doc, search_word)
        
        is_found = len(found_instances) > 0
        if is_found:
            results["found_count"] += 1
            found_terms.append({
                'term': term,
                'instances': found_instances
            })
            print(f"    ✅ FOUND {len(found_instances)} instances")
        else:
            print(f"    ❌ NOT FOUND")
        
        results["part1_results"].append({
            "word": search_word,
            "found": is_found,
            "instances_count": len(found_instances)
        })
        
        results["debug_info"].append(f"Searched '{search_word}': {'FOUND' if is_found else 'NOT FOUND'} ({len(found_instances)} instances)")
    
    print(f"📊 PART 1 Complete: Found {results['found_count']}/{results['total_terms']} words")
    
    # PART 2: Check formatting (only if all words found)
    if results["found_count"] == results["total_terms"]:
        print("🎨 PART 2: Checking formatting...")
        for i, found_term in enumerate(found_terms):
            term = found_term['term']
            instances = found_term['instances']
            
            print(f"  🎨 Checking format for: '{term['searchedWord']}' - {term['formatCheck']} ({i+1}/{len(found_terms)})")
            
            if instances:
                # Check formatting of first instance
                first_instance = instances[0]
                run = first_instance['run']
                paragraph = first_instance['paragraph']
                
                format_result = check_specific_formatting(term, run, paragraph)
                
                if format_result["correct"]:
                    results["format_correct_count"] += 1
                    print(f"    ✅ FORMAT CORRECT")
                else:
                    print(f"    ❌ FORMAT INCORRECT: {format_result['message']}")
                
                results["part2_results"].append({
                    "word": term["searchedWord"],
                    "format_check": term["formatCheck"],
                    "correct": format_result["correct"],
                    "message": format_result["message"],
                    "debug": format_result["debug"]
                })
        
        print(f"📊 PART 2 Complete: {results['format_correct_count']}/{len(found_terms)} correctly formatted")
    else:
        print("⚠️ PART 2 Skipped: Not all words found")
    
    # Calculate overall score
    part1_percentage = (results["found_count"] / results["total_terms"]) * 100
    part2_percentage = (results["format_correct_count"] / results["total_terms"]) * 100 if results["found_count"] == results["total_terms"] else 0
    results["overall_score"] = (part1_percentage + part2_percentage) / 2
    
    print(f"🏆 Final Score: {results['overall_score']:.1f}% (Part1: {part1_percentage:.1f}%, Part2: {part2_percentage:.1f}%)")
    
    return results

def create_comprehensive_practice_file():
    """Create comprehensive practice document with all formatting tasks"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("Comprehensive Formatting Practice Document")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(16)
        
        doc.add_paragraph("")
        doc.add_paragraph("Instructions: Apply the correct formatting to each term below:")
        doc.add_paragraph("")
        
        # Add all terms without formatting (students must add formatting)
        for term in FORMATTING_TERMS:
            para = doc.add_paragraph(term["searchedWord"])
            # Ensure no formatting is applied initially
            para.runs[0].bold = False
            para.runs[0].italic = False
        
        doc.add_paragraph("")
        doc.add_paragraph("Formatting Guide:")
        doc.add_paragraph("• Bold1 - Make bold")
        doc.add_paragraph("• Italic1 - Make italic")
        doc.add_paragraph("• Font sizes: Use Home → Font Size")
        doc.add_paragraph("• Colors: Use Home → Font Color")
        doc.add_paragraph("• Alignment: Use Home → Paragraph alignment")
        doc.add_paragraph("• Special effects: Use Home → Font dialog (Ctrl+D)")
        
        # Save to memory
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Upload to Azure
        blob_client = blob_service_client.get_blob_client(
            container=CONTAINER_NAME,
            blob="comprehensive-formatting-practice.docx"
        )
        
        blob_client.upload_blob(doc_stream.getvalue(), overwrite=True)
        print("✅ Comprehensive practice file uploaded to Azure!")
        return True
        
    except Exception as e:
        print(f"❌ Error creating comprehensive practice file: {e}")
        return False

def get_comprehensive_practice_file_url():
    """Generate download URL for comprehensive practice file"""
    try:
        sas_token = generate_blob_sas(
            account_name=AZURE_ACCOUNT_NAME,
            container_name=CONTAINER_NAME,
            blob_name="comprehensive-formatting-practice.docx",
            account_key=AZURE_ACCOUNT_KEY,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=1)
        )
        
        download_url = f"https://{AZURE_ACCOUNT_NAME}.blob.core.windows.net/{CONTAINER_NAME}/comprehensive-formatting-practice.docx?{sas_token}"
        return download_url
        
    except Exception as e:
        print(f"Error generating download URL: {e}")
        return None

# Enhanced HTML Template with Progress Bars
HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>Comprehensive Word Formatting Checker</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 900px;
            margin: 30px auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .progress-section {
            margin-bottom: 20px;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 8px;
            border-left: 4px solid #007bff;
        }
        .progress-bar {
            background-color: #FF0000;
            border-radius: 8px;
            width: 100%;
            height: 20px;
            margin: 10px 0;
        }
        .progress-fill {
            background-color: #4CAF50;
            height: 100%;
            width: 0%;
            border-radius: 8px;
            transition: width 0.3s ease;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-weight: bold;
            font-size: 12px;
        }
        .step {
            background: #f8f9fa;
            margin: 20px 0;
            padding: 20px;
            border-radius: 8px;
            border-left: 4px solid #007bff;
        }
        button {
            background: #007bff;
            color: white;
            padding: 12px 24px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 16px;
            margin: 10px 5px;
        }
        button:hover {
            background: #0056b3;
        }
        .upload-area {
            border: 2px dashed #ccc;
            padding: 40px;
            text-align: center;
            border-radius: 8px;
            margin: 20px 0;
        }
        .results {
            margin-top: 20px;
            padding: 20px;
            border-radius: 8px;
        }
        .success {
            background: #d4edda;
            border: 1px solid #c3e6cb;
            color: #155724;
        }
        .error {
            background: #f8d7da;
            border: 1px solid #f5c6cb;
            color: #721c24;
        }
        .current-file {
            background: #e3f2fd;
            border: 2px solid #2196f3;
            border-radius: 8px;
            padding: 20px;
            margin: 10px 0;
        }
        .recheck-btn {
            background: #28a745 !important;
        }
        .recheck-btn:hover {
            background: #218838 !important;
        }
        .detail-list {
            font-family: monospace;
            font-size: 12px;
            max-height: 300px;
            overflow-y: auto;
            background: #f8f9fa;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📝 Comprehensive Word Formatting Checker</h1>
        <p>Test your Word document formatting skills with 30+ formatting challenges!</p>
        
        <div class="step">
            <h3>Step 1: Download Practice File</h3>
            <p>Download the comprehensive practice document with all formatting challenges.</p>
            <button onclick="downloadPracticeFile()">📁 Download Comprehensive Practice File</button>
        </div>
        
        <div class="step">
            <h3>Step 2: Apply ALL Required Formatting</h3>
            <p><strong>This is a comprehensive test covering:</strong></p>
            <ul>
                <li><strong>Text Effects:</strong> Bold, Italic, Underline (single, double, dotted), Strikethrough</li>
                <li><strong>Advanced Text:</strong> Superscript, Subscript, Small Caps</li>
                <li><strong>Font Sizes:</strong> 10pt, 12pt, 15pt, 26pt</li>
                <li><strong>Font Types:</strong> Arial, Bookman, Comic Sans, Impact, Tahoma, Verdana</li>
                <li><strong>Colors:</strong> Green text, Dark blue text, Red underlines</li>
                <li><strong>Highlights:</strong> Turquoise highlighting</li>
                <li><strong>Spacing:</strong> Expanded and condensed character spacing</li>
                <li><strong>Alignment:</strong> Left, Center, Right paragraph alignment</li>
            </ul>
            <p><strong>💡 Tips:</strong></p>
            <ul>
                <li>Use <strong>Home → Font Dialog (Ctrl+D)</strong> for advanced formatting</li>
                <li>Use <strong>Home → Paragraph</strong> for alignment</li>
                <li>Character spacing is in <strong>Font Dialog → Advanced tab</strong></li>
            </ul>
        </div>
        
        <div class="step">
            <h3>Step 3: Upload for Comprehensive Analysis</h3>
            <div class="upload-area" id="uploadArea">
                <input type="file" id="fileInput" accept=".docx" style="display: none;" onchange="uploadFile()">
                <button onclick="document.getElementById('fileInput').click()">📤 Upload Document</button>
                <p>Select your completed .docx file</p>
            </div>
            <div id="results"></div>
        </div>
    </div>
    
    <script>
        // Store the uploaded file in memory
        let uploadedFile = null;
        let uploadedFileName = null;

        function downloadPracticeFile() {
            fetch('/download-comprehensive-practice')
                .then(response => response.json())
                .then(data => {
                    if (data.download_url) {
                        window.open(data.download_url, '_blank');
                        alert('Comprehensive practice file download started! Check your downloads folder.');
                    } else {
                        alert('Error: ' + (data.error || 'Could not generate download link'));
                    }
                })
                .catch(error => {
                    alert('Error downloading file: ' + error);
                });
        }
        
        function uploadFile() {
            const fileInput = document.getElementById('fileInput');
            const file = fileInput.files[0];
            
            if (!file) {
                alert('Please select a file');
                return;
            }
            
            if (!file.name.endsWith('.docx')) {
                alert('Please select a .docx file');
                return;
            }
            
            uploadedFile = file;
            uploadedFileName = file.name;
            processUpload(file);
        }
        
        function recheckFile() {
            if (!uploadedFile) {
                alert('No file to recheck. Please upload a file first.');
                return;
            }
            
            const fileInput = document.getElementById('fileInput');
            fileInput.click();
            
            const resultsDiv = document.getElementById('results');
            resultsDiv.innerHTML = `
                <div style="text-align: center; padding: 30px; background: #fff3cd; border-radius: 8px; border: 2px solid #ffc107;">
                    <h3 style="color: #856404; margin-top: 0;">📁 Please Re-select Your File</h3>
                    <p>Since you've made changes to <strong>${uploadedFileName}</strong>, please re-select it from the file dialog.</p>
                    <p><em>This ensures we analyze the latest version of your document.</em></p>
                </div>
            `;
        }
        
        function selectNewFile() {
            uploadedFile = null;
            uploadedFileName = null;
            
            const uploadArea = document.getElementById('uploadArea');
            uploadArea.innerHTML = `
                <input type="file" id="fileInput" accept=".docx" style="display: none;" onchange="uploadFile()">
                <button onclick="document.getElementById('fileInput').click()">📤 Upload Document</button>
                <p>Select your completed .docx file</p>
            `;
            
            document.getElementById('results').innerHTML = '';
            document.getElementById('fileInput').click();
        }
        
        function processUpload(file) {
            console.log('🚀 Starting processUpload with file:', file.name);
            
            const resultsDiv = document.getElementById('results');
            resultsDiv.innerHTML = `
                <div style="text-align: center; padding: 30px; background: #e7f3ff; border-radius: 8px; border: 2px solid #007bff;">
                    <h3 style="color: #007bff; margin-top: 0;">🔍 Comprehensive Analysis in Progress...</h3>
                    <p>Analyzing 30+ formatting elements in: <strong>${file.name}</strong></p>
                    <div style="margin: 20px 0;">
                        <div style="display: inline-block; width: 50px; height: 4px; background: #007bff; border-radius: 2px; animation: pulse 1.5s infinite;"></div>
                    </div>
                </div>
                <style>
                    @keyframes pulse {
                        0% { opacity: 1; }
                        50% { opacity: 0.3; }
                        100% { opacity: 1; }
                    }
                </style>
            `;
            
            const formData = new FormData();
            formData.append('file', file);
            
            console.log('📤 Sending request to /check-comprehensive-formatting');
            console.log('📄 File details:', {
                name: file.name,
                size: file.size,
                type: file.type,
                lastModified: new Date(file.lastModified)
            });
            
            fetch('/check-comprehensive-formatting', {
                method: 'POST',
                body: formData
            })
            .then(response => {
                console.log('📥 Response received:', response.status, response.statusText);
                if (!response.ok) {
                    throw new Error(`HTTP error! status: ${response.status}`);
                }
                return response.json();
            })
            .then(data => {
                console.log('✅ Data received:', data);
                displayComprehensiveResults(data);
                updateUploadButton();
            })
            .catch(error => {
                console.error('💥 Fetch error:', error);
                resultsDiv.innerHTML = 
                    `<div class="results error">
                        <p>❌ Error: ${error.message}</p>
                        <p>Check the browser console (F12) and Flask terminal for more details.</p>
                        <button onclick="location.reload()" style="margin-top: 10px;">🔄 Refresh Page</button>
                    </div>`;
            });
        }
        
        function updateUploadButton() {
            const uploadArea = document.getElementById('uploadArea');
            
            if (uploadedFile) {
                uploadArea.innerHTML = \`
                    <div class="current-file">
                        <p><strong>📄 Current file:</strong> \${uploadedFileName}</p>
                        <div style="margin-top: 15px;">
                            <button onclick="recheckFile()" class="recheck-btn">
                                🔄 Recheck File (Select Updated Version)
                            </button>
                            <button onclick="selectNewFile()" style="background: #6c757d; font-size: 14px; padding: 8px 16px;">
                                📁 Select Different File
                            </button>
                        </div>
                        <p style="font-size: 12px; color: #666; margin-top: 10px;">
                            💡 After editing in Word, save and click "Recheck" to re-select the saved file
                        </p>
                        <input type="file" id="fileInput" accept=".docx" style="display: none;" onchange="uploadFile()">
                    </div>
                `;
            }
        }
        
        function displayComprehensiveResults(data) {
            const resultsDiv = document.getElementById('results');
            const resultClass = data.overall_score >= 90 ? 'success' : 'error';
            
            resultsDiv.innerHTML = '';
            
            let html = `<div class="results ${resultClass}">`;
            
            // Header with timestamp
            html += `<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 15px; border-radius: 8px; margin: -20px -20px 20px -20px;">`;
            html += `<h3 style="margin: 0;">📊 Comprehensive Analysis Complete!</h3>`;
            if (data.timestamp) {
                const timestamp = new Date(data.timestamp).toLocaleTimeString();
                html += `<p style="margin: 5px 0 0 0; opacity: 0.9;">Analyzed at: ${timestamp}</p>`;
            }
            if (data.filename) {
                html += `<p style="margin: 5px 0 0 0; opacity: 0.9;">File: ${data.filename}</p>`;
            }
            html += `</div>`;
            
            // Progress Bars
            html += `<div class="progress-section">`;
            html += `<h4 style="margin: 5px 0;">Overall Progress</h4>`;
            html += `<div class="progress-bar">`;
            html += `<div class="progress-fill" style="width: ${data.overall_score}%">${Math.round(data.overall_score)}%</div>`;
            html += `</div></div>`;
            
            html += `<div class="progress-section">`;
            html += `<h4 style="margin: 5px 0;">Part 1: Finding Words (${data.found_count}/${data.total_terms})</h4>`;
            const part1Percentage = (data.found_count / data.total_terms) * 100;
            html += `<div class="progress-bar">`;
            html += `<div class="progress-fill" style="width: ${part1Percentage}%">${Math.round(part1Percentage)}%</div>`;
            html += `</div></div>`;
            
            if (data.found_count === data.total_terms) {
                html += `<div class="progress-section">`;
                html += `<h4 style="margin: 5px 0;">Part 2: Checking Formatting (${data.format_correct_count}/${data.total_terms})</h4>`;
                const part2Percentage = (data.format_correct_count / data.total_terms) * 100;
                html += `<div class="progress-bar">`;
                html += `<div class="progress-fill" style="width: ${part2Percentage}%">${Math.round(part2Percentage)}%</div>`;
                html += `</div></div>`;
            } else {
                html += `<div class="progress-section">`;
                html += `<h4 style="margin: 5px 0;">Part 2: Checking Formatting</h4>`;
                html += `<div class="progress-bar">`;
                html += `<div style="background-color: #FF0000; height: 100%; width: 100%; border-radius: 8px; display: flex; align-items: center; justify-content: center; color: white; font-weight: bold; font-size: 12px;">`;
                html += `Finish Typing All Words First!`;
                html += `</div></div></div>`;
            }
            
            // Overall Score
            html += `<div style="text-align: center; font-size: 24px; margin: 20px 0; padding: 15px; background: ${data.overall_score >= 90 ? '#d4edda' : data.overall_score >= 70 ? '#fff3cd' : '#f8d7da'}; border-radius: 8px;">`;
            html += `<strong>Overall Score: ${Math.round(data.overall_score)}%</strong>`;
            html += `</div>`;
            
            // Grade
            let grade = 'F';
            if (data.overall_score >= 90) grade = 'A';
            else if (data.overall_score >= 80) grade = 'B';
            else if (data.overall_score >= 70) grade = 'C';
            else if (data.overall_score >= 60) grade = 'D';
            
            html += `<div style="text-align: center; font-size: 32px; margin: 20px 0; padding: 20px; background: ${grade === 'A' ? '#d4edda' : grade === 'F' ? '#f8d7da' : '#fff3cd'}; border-radius: 8px;">`;
            html += `<strong>Grade: ${grade}</strong>`;
            html += `</div>`;
            
            // Part 1 Results (Words Found)
            html += `<details style="margin: 20px 0;"><summary style="cursor: pointer; font-weight: bold; color: #007bff;">📝 Part 1: Word Finding Results (${data.found_count}/${data.total_terms} found)</summary>`;
            html += `<div class="detail-list">`;
            data.part1_results.forEach(result => {
                const icon = result.found ? '✅' : '❌';
                html += `<div style="margin: 2px 0;">${icon} ${result.word} ${result.found ? 'FOUND' : 'NOT FOUND'}</div>`;
            });
            html += `</div></details>`;
            
            // Part 2 Results (Formatting Check)
            if (data.part2_results && data.part2_results.length > 0) {
                html += `<details style="margin: 20px 0;"><summary style="cursor: pointer; font-weight: bold; color: #007bff;">🎨 Part 2: Formatting Results (${data.format_correct_count}/${data.total_terms} correct)</summary>`;
                html += `<div class="detail-list">`;
                data.part2_results.forEach(result => {
                    const icon = result.correct ? '✅' : '❌';
                    const color = result.correct ? 'green' : 'red';
                    html += `<div style="margin: 4px 0; color: ${color};">`;
                    html += `${icon} <strong>${result.word}</strong> - ${result.message}`;
                    html += `<br><span style="margin-left: 20px; font-size: 10px; color: #666;">Debug: ${result.debug}</span>`;
                    html += `</div>`;
                });
                html += `</div></details>`;
            }
            
            // Recheck section
            html += `<div style="text-align: center; margin-top: 20px; padding-top: 15px; border-top: 2px solid #007bff;">`;
            html += `<h4 style="color: #007bff; margin-bottom: 15px;">📤 Need to Make Changes?</h4>`;
            html += `<p>Edit your document in Word, save, and recheck for fresh analysis!</p>`;
            html += `<button onclick="recheckFile()" class="recheck-btn" style="font-size: 16px; padding: 12px 30px;">`;
            html += `🔄 Recheck File (Re-select to Analyze)`;
            html += `</button>`;
            html += `<button onclick="selectNewFile()" style="background: #6c757d; font-size: 14px; padding: 10px 20px; margin-left: 10px;">`;
            html += `📁 Upload Different File`;
            html += `</button>`;
            html += `</div>`;
            
            // Debug information
            if (data.debug_info && data.debug_info.length > 0) {
                html += `<details style="margin-top: 20px;"><summary style="cursor: pointer; color: #007bff;">🔍 Technical Debug Info (click to expand)</summary>`;
                html += `<div class="detail-list">`;
                data.debug_info.forEach(info => {
                    html += `<div style="margin: 2px 0;">${info}</div>`;
                });
                html += `</div></details>`;
            }
            
            html += `</div>`;
            resultsDiv.innerHTML = html;
            
            // Scroll to results
            resultsDiv.scrollIntoView({ behavior: 'smooth', block: 'start' });
        }
    </script>
</body>
</html>
"""

# Routes
@app.route('/')
def home():
    """Main page"""
    return HTML_TEMPLATE

@app.route('/setup')
def setup():
    """Setup route - run this first to create and upload comprehensive practice file"""
    try:
        if create_comprehensive_practice_file():
            return "✅ Setup complete! Comprehensive practice file uploaded to Azure. Go to <a href='/'>home page</a> to test."
        else:
            return "❌ Setup failed. Check your Azure credentials."
    except Exception as e:
        return f"❌ Setup error: {str(e)}"

@app.route('/download-comprehensive-practice')
def download_comprehensive_practice():
    """Provide download URL for comprehensive practice file"""
    try:
        download_url = get_comprehensive_practice_file_url()
        if download_url:
            return jsonify({"download_url": download_url})
        else:
            return jsonify({"error": "Could not generate download URL"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/check-comprehensive-formatting', methods=['POST'])
def check_comprehensive_formatting_route():
    """Check comprehensive formatting in uploaded document"""
    print("📥 Route called: /check-comprehensive-formatting")
    
    try:
        if 'file' not in request.files:
            print("❌ No file in request")
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        print(f"📄 File received: {file.filename}")
        
        if file.filename == '':
            print("❌ Empty filename")
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.endswith('.docx'):
            print(f"❌ Wrong file type: {file.filename}")
            return jsonify({'error': 'Please upload a .docx file'}), 400
        
        print("📖 Reading file content...")
        # Read the uploaded file
        file_content = file.read()
        print(f"📊 File size: {len(file_content)} bytes")
        
        if len(file_content) == 0:
            print("❌ File is empty")
            return jsonify({'error': 'Uploaded file is empty'}), 400
        
        file_stream = io.BytesIO(file_content)
        
        print("🔍 Creating Document object...")
        doc = Document(file_stream)
        print(f"📄 Document loaded. Paragraphs: {len(doc.paragraphs)}")
        
        print("🚀 Starting comprehensive analysis...")
        # Run comprehensive analysis
        results = check_comprehensive_formatting(doc)
        print(f"✅ Analysis complete. Found: {results['found_count']}/{results['total_terms']}")
        
        # Format response
        response = {
            "total_terms": results["total_terms"],
            "found_count": results["found_count"],
            "format_correct_count": results["format_correct_count"],
            "overall_score": results["overall_score"],
            "part1_results": results["part1_results"],
            "part2_results": results["part2_results"],
            "debug_info": results["debug_info"],
            "timestamp": datetime.now().isoformat(),
            "filename": file.filename
        }
        
        print("📤 Sending response...")
        # Add cache-busting headers
        from flask import make_response
        resp = make_response(jsonify(response))
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        
        return resp
        
    except Exception as e:
        print(f"💥 Error in route: {str(e)}")
        import traceback
        traceback.print_exc()
        
        error_response = {
            'error': 'Failed to analyze document',
            'message': str(e),
            'overall_score': 0,
            'total_terms': len(FORMATTING_TERMS),
            'found_count': 0,
            'format_correct_count': 0,
            'timestamp': datetime.now().isoformat(),
            'filename': file.filename if 'file' in locals() and file else 'unknown'
        }
        
        from flask import make_response
        resp = make_response(jsonify(error_response), 500)
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        
        return resp

if __name__ == '__main__':
    print("🚀 Starting Comprehensive Word Formatting Checker...")
    print("📝 First visit: http://localhost:5000/setup")
    print("📝 Then visit: http://localhost:5000")
    print("⚙️  Make sure to set Azure environment variables!")
    app.run(debug=True, host='0.0.0.0', port=5000)